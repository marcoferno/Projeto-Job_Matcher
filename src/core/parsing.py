"""
    Realiza a leitura e extração de texto de currículos, suportando arquivos .pdf, .docx e .txt,
    com fallbacks opcionais (pdfminer/OCR) para PDFs sem texto e normalização de whitespace para melhorar
    o ranking nas buscas.
"""

# IMPORTS (Recurso de Linguagem/Compatibilidade)
from __future__ import annotations

# IMPORTS (Stdlib)
from dataclasses import dataclass
from pathlib import Path
from typing import Literal, Optional
import docx
import PyPDF2
import pytesseract
import re # Usado como fallback para limpar HTML quando bs4 não estiver disponível.


# IMPORTS (Terceiros)
try:
    from loguru import logger
except Exception:
    class _Dummy:
        def info(self, *a, **k): pass
        def warning(self, *a, **k): pass
        def error(self, *a, **k): pass
        def debug(self, *a, **k): pass
    logger = _Dummy()

# ======================================================================
# Opções e erros
# ======================================================================

@dataclass(slots = True)
class ExtractOptions:
    """
        Opções avançadas de extração de texto.

        Controla qual backend usar para PDF (PyPDF2/OCR), se o texto
        deve ser normalizado e o tamanho máximo de texto retornado.
    """
    prefer_pdfminer: bool = False
    ocr_if_scanned: bool = False
    ocr_langs: tuple[str, ...] = ("pt", "en")
    normalize_whitespace: bool = True
    max_chars: int = 600_000

class ExtractionError(Exception): ...
""" Erro genérico de extração de Texto """

class UnsupportedFormatError(ExtractionError): ...
"""Erro para extensões/formatos de arquivo não suportados."""

# ======================================================================
# Implementações de leitura
# ======================================================================

def _read_txt(path: Path) -> str:
    """
        Lê arquivo .txt tentando primeiro UTF-8 e depois Latin-1.
        Também remove BOM (se existir) e normaliza as quebras de linha para o formato '\n'.
    """
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="latin-1", errors="strict")
    return text.lstrip("\ufeff").replace("\r\n", "\n").replace("\r", "\n")

def _read_pdf_pypdf2(path: Path) -> str:
    """
        Extrai texto de um PDF usando PyPDF2.

        Tenta descriptografar PDFs protegidos com senha vazia.
        Caso falhe ou ocorram erros na extração de alguma página, essas páginas são simplesmente ignoradas.
    """
    text_parts: list[str] = []

    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        if getattr(reader, "is_encrypted", False):
            try:
                reader.decrypt("")
            except Exception:
                logger.error("PDF criptografado e sem senha: %s", path)
                return ""
        for page in reader.pages:
            try:
                text_parts.append(page.extract_text() or "")
            except Exception:
                continue
    return "\n".join(text_parts)

def _read_pdf_pdfminer(path: Path) -> str:

    try:
        from pdfminer.high_level import extract_text
    except Exception as e: # Pdfminer não instalado ou com problema ≥ Avisa no log e retorna vazio
        logger.debug(f"pdfminer não disponível: {e}")
        return ""

    try:
        return extract_text(str(path)) or ""
    except Exception as e:
        logger.error(f"Falha ao ler PDF com pdfminer: {path} - {e}")
        return ""

def _read_pdf_ocr(path: Path, langs: tuple[str, ...]) -> str:
    """
        Para OCR real de PDF, ideal usar pdf2image + pytesseract/easyocr.
        Manter opcional para não exigir dependências extras.
    """
    try:
        from pdf2image import convert_from_path  # Requer poppler instalado no SO (Instalar depois)
    except Exception as e:
        logger.warning(f"OCR desabilitado (dependências ausentes): {e}")
        return ""

    try:
        pages = convert_from_path(str(path), dpi=200)
    except Exception as e:
        logger.error(f"Falha ao converter PDF em imagens para OCR: {path} - {e}")
        return ""

    text_parts: list[str] = []

    try:
        langs_str = "+".join(langs) if langs else "por+eng"
        for img in pages:
            txt = pytesseract.image_to_string(img, lang=langs_str) or ""
            text_parts.append(txt)
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Falha ao aplicar OCR no PDF: {path} - {e}")
        return ""

def _read_docx(path: Path) -> str:
    """
        Extrai texto de um arquivo .docx (Word).

        Lê tanto os parágrafos principais quanto textos presentes em tabelas nos currículos.
    """
    doc = docx.Document(str(path))
    parts: list[str] = []
    parts.extend(p.text for p in doc.paragraphs if p.text)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                cell_text = "\n".join(p.text for p in cell.paragraphs if p.text)
                if cell_text:
                    parts.append(cell_text)
    return "\n".join(parts)

# ======================================================================
# Orquestração
# ======================================================================

def _normalize_text(text: str, max_chars: int, do_ws: bool) -> str:
    """
        Aplica cortes e normalização leve ao texto extraído.

        Limitando o texto ao máximo de aracteres e opcionalmente colapsa múltiplos espaços num só.
        Reduzindo blocos de muitas quebras de linha para no máximo duas.
    """
    text = text[:max_chars]
    if do_ws:
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        text = text.strip()
    return text

def extrair_texto_adv(arquivo: str | Path, *, options: Optional[ExtractOptions] = None) -> str:
    """
        Extrai texto com opções avançadas (pdfminer/ocr/normalização).
        Levanta FileNotFoundError para caminho inexistente e UnsupportedFormatError para extensão inválida.
    """
    options = options or ExtractOptions()
    path = Path(arquivo)
    if not path.exists():
        raise FileNotFoundError(f"Arquivo não encontrado: {arquivo}")

    ext = path.suffix.lower()

    if ext not in {".pdf", ".docx", ".txt"}:
        raise UnsupportedFormatError(f"Extensão não suportada: {ext}")
    if ext == ".txt":
        text = _read_txt(path)
    elif ext == ".docx":
        text = _read_docx(path)
    else:
        text = ""
        if options.prefer_pdfminer:
            text = _read_pdf_pdfminer(path)
            if not text:
                logger.info("pdfminer sem texto; tentando PyPDF2: %s", path)
                text = _read_pdf_pypdf2(path)
        else:
            text = _read_pdf_pypdf2(path)
            if not text:
                logger.info("PyPDF2 sem texto; tentando pdfminer: %s", path)
                text = _read_pdf_pdfminer(path)
            if not text and options.ocr_if_scanned:
                logger.info("PDF possivelmente scaneado; tentando OCR: %s", path)
                text = _read_pdf_ocr(path, options.ocr_langs)
    return _normalize_text(text, options.max_chars, options.normalize_whitespace)

def extrair_texto(arquivo: str) -> str:
    """
        Versão compatível com o MVP.

        Usa o fluxo avançado de extração, mas sem OCR, e mantém compatibilidade com o comportamento antigo:
        se a extensão não for reconhecida, tenta ler o arquivo como .txt.
    """
    try:
        return extrair_texto_adv(
            arquivo,
            options=ExtractOptions(
                ocr_if_scanned=True,
                prefer_pdfminer=False,
            ),
        )
    except UnsupportedFormatError:
        return _read_txt(Path(arquivo))